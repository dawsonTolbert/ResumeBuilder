from docx import Document
from docx.shared import Inches, Cm

def createResume(projectsFile, infoFile):
    with open(projectsFile, 'r') as file:
        education = file.readline()
        projects = file.readlines()
    file.close()

    with open(infoFile, 'r') as file:
        info = file.readline()
    file.close()

    document = Document()

    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    personalInfo = info.strip().split(';')
    name, number, email, address = personalInfo
    document.add_heading(name, level=1)
    document.add_paragraph(f'{address}, {email}, {number}')

    edu_info = education.strip().split(';')
    if len(edu_info) == 5:
        degree, school, location, date, gpa = edu_info
        document.add_heading(f'Education', level=2)
        document.add_paragraph(f'{school}, {location}')
        document.add_paragraph(f'{degree}, {date}. GPA: {gpa}.')
    else:
        print(f"Issue processing education")

    for project in projects:
        project_info = project.strip().split(';')
        if len(project_info) <= 1:
            print(f"Issue processing project: {project}")
        else:
            project_name = project_info[0]
            document.add_heading(f'Project: {project_name}', level=2)
            for bullet in project_info[1:]:
                document.add_paragraph(f'{bullet}', style='List Bullet')

    document.save('Resume.docx')
    print("Resume created successfully.")


def main():
    projectsFile = 'projects.txt'
    infoFile = 'info.txt'
    createResume(projectsFile, infoFile)


if __name__ == "__main__":
    main()
